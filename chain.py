
import os
from langchain import LangChain
from docx import Document

# Read all the files in the prompt folder

prompt_folder = 'prompt_folder/'
prompt_list = []
for file in os.listdir(prompt_folder):
    with open(prompt_folder + file, 'r') as f:
        context = f.read()
        prompt_list.append(context)

# Read all the files in the code folder
code_folder = 'code_folder/'
code_list = []


# Create a new Document
        
def create_docx(file_name, context_list):
    document = Document()

    # Add a Title to the document
    document.add_heading('Code Explanation', 0)

    # Add a paragraph to the document
    for context in context_list:
      document.add_paragraph(
          context
      )

      document.add_page_break()   

    documet_name = 'output_folder/' + file_name + '.docx'
    document.save(documet_name)


    # create lang chain for each prompt in prompt_list and pass the code form code_list as input to first prompt

for file in os.listdir(code_folder):
    with open(code_folder + file, 'r') as f:
        code = f.read()

        lang_chain = LangChain(
            model="gpt-4",
            temperature=0.7,
            max_tokens=1024,
            api_key='sk-XXXXX'
        )

        document_context_list = []

        model_input = code
        lang_chain.set_code(model_input)
        for prompt in prompt_list:
            lang_chain.set_prompt(prompt)
            response = lang_chain.generate_response()
            document_context_list.append(response)
            lang_chain.set_code(response)

        create_docx(file.split('.')[0], document_context_list)
