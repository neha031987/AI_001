from openai import OpenAI

from docx import Document
# from docx.shared import Inches
import os


prompt_folder = 'prompt_folder/'
prompt_list = []

code_folder = 'code_folder/'

# Read all the files in the prompt folder
# and append the content to the prompt_list
for file in os.listdir(prompt_folder):
    with open(prompt_folder + file, 'r') as f:
        context = f.read()
        prompt_list.append(context)

# Create a new Document
def create_docx(file_name, response_list):
    document = Document()

    # Add a Title to the document
    document.add_heading('Code Explanation', 0)

    # Add a paragraph to the document
    for resp in response_list:
      document.add_paragraph(
          resp
      )

      document.add_page_break()   

    documet_name = 'output_folder/' + file_name + '.docx'
    document.save(documet_name)

client = OpenAI()
# we can set the api key either here or in the .env file.
client.set_api_key('sk-XXXXX')
# openAI secret 

# Create a new document for each prompt and code

def call_openai(prompt, code_file_context):
  response = client.chat.completions.create(
    model="gpt-4",
    messages=[
      {
        "role": "system",
        "content": prompt
      },
      {
        "role": "user",
        "content": code_file_context
      }
    ],
    temperature=0.7,
    max_tokens=1024,
    top_p=1
  )

  return response

# Read all the files in the code folder
# and append the content to the code_list

if __name__ == '__main__':
   
  for file in os.listdir(code_folder):
      with open(code_folder + file, 'r') as f:
          code_file_context = f.read()

          response_list = []
          for prompt in prompt_list:
            response = call_openai(prompt, code_file_context)
            response_list.append(response['choices'][0]['message']['content'])

          # create a new document for each prompt and code
            # pass the file naem without extension for the word document name
          create_docx(file.split('.')[0], response_list)
